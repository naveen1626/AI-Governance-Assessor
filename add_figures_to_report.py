"""
Add generated figures to the final report document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Paths
report_path = "D:/AI/Projects/AI_Governance/AI_Dual_Use_Risk_Assessor_Report.docx"
figures_dir = Path("D:/AI/Projects/AI_Governance/report_figures")
output_path = "D:/AI/Projects/AI_Governance/AI_Dual_Use_Risk_Assessor_Report_Final.docx"

# Load document
doc = Document(report_path)

# Figure definitions with captions and target sections
figures = [
    {
        "file": "fig1_risk_distribution.png",
        "caption": "Figure 1: Risk Tier Distribution across 74 assessed research papers. The majority (56.8%) received Medium tier classification, while 43.2% were classified as High or Critical risk.",
        "after_section": "4. Results"
    },
    {
        "file": "fig2_category_distribution.png",
        "caption": "Figure 2: Distribution of assessed papers by research category. AI/ML papers were most frequent (17), followed by Semiconductor and Biomedical (13 each).",
        "after_section": "4. Results"
    },
    {
        "file": "fig3_axis_radar.png",
        "caption": "Figure 3: Average scores across the 12-axis risk framework. The radar chart shows C2 (Mitigation Plans) as the most stressed axis with an average of 2.0, indicating a common gap in deployment safeguards.",
        "after_section": "4. Results"
    },
    {
        "file": "fig4_section_averages.png",
        "caption": "Figure 4: Average risk scores by framework section. Section C (Safeguards & Governance) shows the highest average score (1.77), indicating this area requires the most attention across assessed papers.",
        "after_section": "4. Results"
    },
    {
        "file": "fig5_capability_safeguard.png",
        "caption": "Figure 5: Capability Index vs Safeguard Index comparison. The negative governance gap (-0.22) indicates that on average, safeguard measures slightly exceed capability risks in the assessed portfolio.",
        "after_section": "4. Results"
    },
    {
        "file": "fig6_summary_dashboard.png",
        "caption": "Figure 6: Evaluation summary dashboard showing key metrics from the assessment of 74 research papers across six dual-use categories.",
        "after_section": "4. Results"
    },
    {
        "file": "fig7_accuracy_table.png",
        "caption": "Table 1: Category detection accuracy by research domain. Overall accuracy of 97.3% (72/74 correct classifications).",
        "after_section": "4. Results"
    }
]

# Find the Results section and add figures after first content paragraph
results_section_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "4. Results":
        results_section_idx = i
        break

if results_section_idx is None:
    print("ERROR: Could not find Results section")
    exit(1)

# Find the end of Results section content (before Discussion)
insert_idx = results_section_idx + 1
for i in range(results_section_idx + 1, len(doc.paragraphs)):
    if doc.paragraphs[i].style and 'Heading' in doc.paragraphs[i].style.name:
        insert_idx = i
        break
    if doc.paragraphs[i].text.strip():
        insert_idx = i + 1

print(f"Found Results section at paragraph {results_section_idx}")
print(f"Will insert figures before paragraph {insert_idx}")

# Insert figures in reverse order (so they appear in correct order)
for fig_info in reversed(figures):
    fig_path = figures_dir / fig_info["file"]

    if not fig_path.exists():
        print(f"WARNING: Figure not found: {fig_path}")
        continue

    # Create new paragraph for figure
    new_para = doc.paragraphs[insert_idx - 1]._element

    # Add figure
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    # Determine width based on figure type
    if "summary_dashboard" in fig_info["file"]:
        width = Inches(6.5)
    elif "radar" in fig_info["file"]:
        width = Inches(5.0)
    else:
        width = Inches(5.5)

    run.add_picture(str(fig_path), width=width)

    # Add caption
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(fig_info["caption"])
    caption_run.font.size = Pt(10)
    caption_run.italic = True

    # Add spacing
    spacing_para = doc.add_paragraph()

    # Move paragraphs to correct position
    new_para.addnext(spacing_para._element)
    new_para.addnext(caption_para._element)
    new_para.addnext(p._element)

    print(f"Added: {fig_info['file']}")

# Save updated document
doc.save(output_path)
print(f"\nReport with figures saved to: {output_path}")
